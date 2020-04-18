from tkinter import *
import pyperclip
from docx import Document
from docx.shared import Pt       #Helps to specify font size
from docx.shared import RGBColor #Helps to specify font Color


def restart():
    global count, all_text, count_label
    count = 0
    all_text = []
    count_label['text'] = "Add Count: " + str(count)


def popup_msg(msg):
    popup = Tk()
    popup.wm_title("!")
    label = Label(popup, text=msg)
    label.pack(side="top", fill="x", pady=10)
    ok_button = Button(popup, text="Okay", command = popup.destroy)
    ok_button.pack()
    popup.mainloop()


def check_equal(merged_list):
    for text in merged_list:
        if len(text) != len(merged_list[0]):
            return FALSE
    return TRUE


def set_input():
    e1.insert(END, "\n\n"+pyperclip.paste())


def clear():
    e1.delete(1.0, END)


def add():
    global count, all_text
    all_text.append(e1.get(1.0, END))
    e1.delete(1.0, END)
    count += 1
    count_label['text'] = "Add Count: "+str(count)


def finnish():
    global all_text
    merged_list = []
    clear()
    for text in all_text:
        merged_list.append(list(filter(None, text.splitlines())))
    if check_equal(merged_list) == TRUE:
        document = Document()
        for i in range(0, len(merged_list[0])):
            p = document.add_paragraph()
            p.line_spacing = 0.5
            for j in range(0, len(merged_list)):
                wp = p.add_run(merged_list[j][i])
                wp.font.size = Pt(13)
                if j == 1:
                    wp.font.color.rgb = RGBColor(139, 69, 19)
                else:
                    wp.font.color.rgb = RGBColor(0, 0, 0)
                wp.add_break()
            wp.add_break()
        document.save('novel.docx')
        restart()
    else:
        restart()
        popup_msg("บรรทัดไม่เท่ากัน โปรดทำใหม่อีกครั้ง")


def copy():
    copy_text = e1.get(1.0, END)
    pyperclip.copy(copy_text)


count = 0
all_text = []

master = Tk()
upperFrame = Frame(master)
upperFrame.pack()

lowerFrame = Frame(master)
lowerFrame.pack(side=BOTTOM)

textLabel = Label(upperFrame, text="Text")
textLabel.pack(side=TOP)

e1 = Text(lowerFrame, height=20, width=40)
e1.pack()
e1.tag_config('green', foreground="green")
e1.tag_config('yellow', background = "yellow")

clear_button = Button(lowerFrame, text="Clear", command=clear)
clear_button.pack(side=LEFT)
paste_button = Button(lowerFrame, text="Paste", command=set_input)
paste_button.pack(side=LEFT)
finnish_button = Button(lowerFrame, text="Finnish", command=finnish)
finnish_button.pack(side=RIGHT)

copy_button = Button(lowerFrame, text="Copy", command=copy)
copy_button.pack(side=BOTTOM)
count_label = Label(lowerFrame, text="Add Count: " + str(count))
count_label.pack(side=BOTTOM)
add_button = Button(lowerFrame, text="Add", command=add)
add_button.pack(side=BOTTOM)
master.mainloop()